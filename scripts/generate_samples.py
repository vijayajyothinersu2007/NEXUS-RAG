"""Generate synthetic enterprise sample documents for Phase 1 parser tests and demos."""

from __future__ import annotations

import sys
from pathlib import Path

import pymupdf
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
SAMPLES = ROOT / "data" / "samples"


POLICY_2025 = """EMPLOYEE POLICY 2025
Document ID: POL-EMP-2025
Version: 2025.1
Owner: Human Resources
Effective: 1 January 2025
References: Security Regulation SR-2024; Compliance Guidelines CG-2025

1. PURPOSE
This policy defines employee attendance, leave, and approval requirements for NexusCorp.

2. ATTENDANCE REQUIREMENT
Employees must maintain at least 70% attendance in each calendar quarter.
Employees below 70% attendance must complete a performance improvement plan with their manager.

3. LEAVE REQUEST PROCESS
Employees must submit leave requests within 7 days of the intended leave date.
Leave requests require approval from the department manager.

4. REMOTE WORK
Remote work is permitted up to 1 day per week with manager approval.

5. SECURITY ALIGNMENT
This policy applies to the employee onboarding process and must remain consistent with Security Regulation SR-2024.
"""

POLICY_2026 = """EMPLOYEE POLICY 2026
Document ID: POL-EMP-2026
Version: 2026.1
Owner: Human Resources
Effective: 1 January 2026
Supersedes: Employee Policy 2025 (POL-EMP-2025)
References: Security Regulation SR-2024; Compliance Guidelines CG-2025

1. PURPOSE
This policy defines employee attendance, leave, and approval requirements for NexusCorp.

2. ATTENDANCE REQUIREMENT
Employees must maintain at least 75% attendance in each calendar quarter.
Employees below 75% attendance must complete a performance improvement plan with HR and their manager.

3. LEAVE REQUEST PROCESS
Employees must submit leave requests within 5 days of the intended leave date.
Leave requests require approval from the department manager and Human Resources.

4. REMOTE WORK
Remote work is permitted up to 3 days per week with manager approval.
Remote workers must complete MFA enrollment before the first remote day.

5. SECURITY ALIGNMENT
This policy applies to the employee onboarding process and must remain consistent with Security Regulation SR-2024.
MFA is required before system access is granted.
"""

SECURITY_REG = """SECURITY REGULATION SR-2024
Document ID: SR-2024
Version: 1.2
Owner: Information Security
Effective: 1 June 2024

1. SCOPE
This regulation affects the employee onboarding process, remote access, and production systems at NexusCorp.

2. ACCESS CONTROL
Multi-factor authentication (MFA) is required for all employees before system access is granted.
Access reviews must be completed quarterly by system owners.

3. ONBOARDING IMPACT
During onboarding, HR must not provision accounts until MFA enrollment is complete.
This regulation affects the leave and remote-work processes defined in Employee Policy documents.

4. REFERENCES
Related documents: Employee Policy 2025, Employee Policy 2026, Operations Manual, Compliance Guidelines CG-2025.
"""

OPS_MANUAL = """OPERATIONS MANUAL
Document ID: OPS-MAN-2025
Version: 3.0
Owner: Operations
Effective: 15 March 2025

1. REQUEST SUBMISSION PROCESS
Employees must submit operational and leave requests within 7 days of the required date.
Requests are routed to the department manager for approval.

2. ONBOARDING PROCESS
The onboarding process requires:
- Identity verification
- Policy acknowledgement
- System access request
- Manager approval

3. CONFLICT NOTE
Request timing in this manual is 7 days. Later policy versions may specify a different window.
This manual references Security Regulation SR-2024 for access provisioning.
"""

COMPLIANCE = """COMPLIANCE GUIDELINES CG-2025
Document ID: CG-2025
Version: 2.0
Owner: Compliance
Effective: 1 April 2025

1. REQUIREMENT TRACEABILITY
Every operational requirement must be supported by a named policy or regulation.
The current employee attendance requirement is defined by the latest Employee Policy.

2. DOCUMENT HIERARCHY
Security Regulation SR-2024 prevails over internal manuals when access-control requirements conflict.
Employee Policy 2026 supersedes Employee Policy 2025.

3. EVIDENCE STANDARD
Compliance reviews must cite document name, version, section, and excerpt. Unsupported claims are not acceptable.
"""


def write_txt() -> None:
    SAMPLES.mkdir(parents=True, exist_ok=True)
    mapping = {
        "employee_policy_2025.txt": POLICY_2025,
        "employee_policy_2026.txt": POLICY_2026,
        "security_regulation_sr2024.txt": SECURITY_REG,
        "operations_manual.txt": OPS_MANUAL,
        "compliance_guidelines.txt": COMPLIANCE,
    }
    for name, content in mapping.items():
        (SAMPLES / name).write_text(content.strip() + "\n", encoding="utf-8")


def write_pdf(name: str, title: str, body: str) -> None:
    path = SAMPLES / name
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 56), title, fontsize=16)
    page.insert_textbox(pymupdf.Rect(72, 88, 540, 760), body, fontsize=11, align=0)
    doc.save(path)
    doc.close()


def write_docx(name: str, title: str, body: str) -> None:
    document = Document()
    heading = document.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for block in body.strip().split("\n\n"):
        lines = block.splitlines()
        if lines and lines[0][:1].isdigit() and "." in lines[0][:4]:
            document.add_heading(lines[0], level=2)
            if len(lines) > 1:
                document.add_paragraph("\n".join(lines[1:]))
        else:
            document.add_paragraph(block)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Field"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Document"
    table.cell(1, 1).text = title
    document.save(SAMPLES / name)


def write_tabular() -> None:
    frame = pd.DataFrame(
        [
            {"requirement": "Attendance minimum", "policy_2025": "70%", "policy_2026": "75%", "status": "modified"},
            {
                "requirement": "Leave request window",
                "policy_2025": "7 days",
                "policy_2026": "5 days",
                "status": "modified",
            },
            {
                "requirement": "Leave approval",
                "policy_2025": "Department manager",
                "policy_2026": "Department manager and HR",
                "status": "modified",
            },
            {"requirement": "MFA before access", "policy_2025": "Not stated", "policy_2026": "Required", "status": "added"},
        ]
    )
    frame.to_csv(SAMPLES / "requirement_matrix.csv", index=False)
    with pd.ExcelWriter(SAMPLES / "requirement_matrix.xlsx", engine="openpyxl") as writer:
        frame.to_excel(writer, sheet_name="Requirements", index=False)


def main() -> int:
    SAMPLES.mkdir(parents=True, exist_ok=True)
    write_txt()
    write_pdf("employee_policy_2025.pdf", "Employee Policy 2025", POLICY_2025)
    write_pdf("employee_policy_2026.pdf", "Employee Policy 2026", POLICY_2026)
    write_pdf("security_regulation_sr2024.pdf", "Security Regulation SR-2024", SECURITY_REG)
    write_docx("operations_manual.docx", "Operations Manual", OPS_MANUAL)
    write_docx("compliance_guidelines.docx", "Compliance Guidelines CG-2025", COMPLIANCE)
    write_tabular()
    print(f"Wrote sample documents to {SAMPLES}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
