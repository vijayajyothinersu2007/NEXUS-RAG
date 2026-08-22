"""Parser tests for PDF, DOCX, TXT, CSV, and XLSX."""

from __future__ import annotations

from pathlib import Path

import pymupdf
import pandas as pd
from docx import Document

from backend.ingestion.parsers import parse_document
from backend.ingestion.parsers.txt_parser import parse_txt


def test_txt_parser_extracts_headings_and_body(tmp_path: Path):
    path = tmp_path / "policy.txt"
    path.write_text("# Attendance\nEmployees must maintain 75% attendance.\n", encoding="utf-8")
    full_text, pages, sections, tables, warnings = parse_txt(path)
    assert "75%" in full_text
    assert pages[0].page_number == 1
    assert any("Attendance" in section.title for section in sections)
    assert tables == []
    assert warnings == []


def test_pdf_parser_keeps_page_numbers(tmp_path: Path):
    path = tmp_path / "manual.pdf"
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Operations Manual\nLeave requests must be submitted within 7 days.")
    page2 = doc.new_page()
    page2.insert_text((72, 72), "Section 2\nMFA is required for all systems.")
    doc.save(path)
    doc.close()

    full_text, pages, _sections, _tables, warnings = parse_document(path)
    assert len(pages) == 2
    assert pages[0].page_number == 1
    assert pages[1].page_number == 2
    assert "7 days" in full_text
    assert "MFA" in full_text
    assert not any("Unable" in warning for warning in warnings)


def test_docx_parser_reads_headings_and_tables(tmp_path: Path):
    path = tmp_path / "policy.docx"
    document = Document()
    document.add_heading("Employee Policy 2026", level=1)
    document.add_paragraph("Employees must submit leave requests within 5 days.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Requirement"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Attendance"
    table.cell(1, 1).text = "75%"
    document.save(path)

    full_text, _pages, sections, tables, _warnings = parse_document(path)
    assert "5 days" in full_text
    assert any("Employee Policy 2026" in section.title for section in sections)
    assert tables
    assert "75%" in tables[0].text_representation


def test_csv_parser_renders_table(tmp_path: Path):
    path = tmp_path / "reqs.csv"
    pd.DataFrame({"requirement": ["MFA"], "owner": ["Security"]}).to_csv(path, index=False)
    full_text, _pages, sections, tables, _warnings = parse_document(path)
    assert tables[0].headers == ["requirement", "owner"]
    assert "MFA" in full_text
    assert sections[0].title == "Sheet1"


def test_xlsx_parser_reads_named_sheet(tmp_path: Path):
    path = tmp_path / "reqs.xlsx"
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        pd.DataFrame({"control": ["Access review"], "frequency": ["Quarterly"]}).to_excel(
            writer, sheet_name="Controls", index=False
        )
    full_text, _pages, sections, tables, _warnings = parse_document(path)
    assert tables[0].sheet_name == "Controls"
    assert "Quarterly" in full_text
    assert sections[0].title == "Controls"


def test_malformed_pdf_does_not_crash_process(tmp_path: Path):
    path = tmp_path / "broken.pdf"
    path.write_bytes(b"%PDF-this-is-broken")
    try:
        parse_document(path)
    except ValueError:
        pass
